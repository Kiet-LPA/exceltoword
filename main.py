import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO
import os
import re
import requests
from openpyxl import load_workbook
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement

st.set_page_config(page_title="Tạo file Word từ Excel", page_icon="📝")

st.title("Tạo file Word từ Excel (sử dụng cột Design Card + STT Dàn Trang)")
sheet_file = st.file_uploader("📁 Tải lên file Excel (.xlsx) có cột 'Design Card' và 'STT Dàn Trang'", type=["xlsx"])
#export_pdf = st.checkbox("Xuất ra PDF (.pdf)")

if st.button("Tạo file Word"):
    if not sheet_file:
        st.warning("⚠️ Vui lòng tải lên file trước.")
    else:
        try:
            excel_bytes = BytesIO(sheet_file.read())
            excel_bytes.seek(0)
            wb = load_workbook(excel_bytes, data_only=True)
            ws = wb.active
            df = pd.read_excel(excel_bytes)

            if 'Design Card' not in df.columns or 'STT Dàn Trang' not in df.columns:
                st.error("❌ Thiếu cột 'Design Card' hoặc 'STT Dàn Trang' trong file Excel.")
            else:
                doc = Document()
                image_map = {}

                col_design = df.columns.get_loc("Design Card") + 1
                col_stt = df.columns.get_loc("STT Dàn Trang") + 1

                for i in range(len(df)):
                    row_excel = i + 2  # Excel index starts at 2
                    cell_link = ws.cell(row=row_excel, column=col_design)
                    raw_stt_cell = ws.cell(row=row_excel, column=col_stt)

                    # Lấy link ảnh
                    if cell_link.hyperlink:
                        image_url = cell_link.hyperlink.target
                    else:
                        image_url = str(cell_link.value).strip() if cell_link.value else None

                    if not image_url:
                        continue

                    # Thêm https nếu thiếu
                    if not image_url.startswith("http"):
                        image_url = "https://" + image_url

                    # Lấy STT
                    stt_raw = str(raw_stt_cell.value).strip() if raw_stt_cell.value else f"Row_{row_excel}"
                    stts = re.split(r"[,\s]+", stt_raw)

                    try:
                        match = re.search(r"drive\.google\.com/file/d/([a-zA-Z0-9_-]+)", image_url)
                        if match:
                            file_id = match.group(1)
                            image_url = f"https://drive.google.com/uc?export=download&id={file_id}"

                        response = requests.get(image_url, timeout=30)
                        response.raise_for_status()
                        img = Image.open(BytesIO(response.content))
                        img.thumbnail((500, 9999))

                        for stt in stts:
                            safe_stt = re.sub(r'[^\w\-_.]', '_', stt)
                            temp_path = f"temp_{safe_stt}.png"
                            img.save(temp_path)
                            image_map[stt] = {
                                "path": temp_path,
                                "valid": True
                            }
                            st.success(f"✅ Ảnh STT {stt} xử lý thành công")

                    except Exception as e:
                        # Ghi log nội bộ nếu cần, nhưng không in ra giao diện
                        print(f"Lỗi khi tải {image_url}: {e}")  # hoặc ghi file log

                        for stt in stts:
                            image_map[stt] = {
                                "path": None,
                                "valid": False
                            }
                            st.warning(f"⚠️ Lỗi ảnh STT {stt}: Link không chứa ảnh")



                if not image_map:
                    st.warning("⚠️ Không có ảnh nào được tải thành công, không thể tạo file.")
                    st.stop()

                # === Dàn ảnh vào Word (6 hình/trang: 3 hàng x 2 cột)
                def sort_key(stt):
                    try:
                        return int(stt)
                    except:
                        return float('inf')

                # Chỉ lấy những ảnh hợp lệ để dàn trang
                sorted_stts = sorted(
                    [stt for stt in image_map if image_map[stt]["valid"]],
                    key=sort_key
                )
                rows_per_page = 3
                cols_per_page = 2
                max_per_page = rows_per_page * cols_per_page
                image_width = Cm(7.83)
                col_img_width = image_width + Cm(0.1)  # hoặc 0.2 nếu vẫn bị khuất

                # Xóa header và footer nếu có
                for section in doc.sections:
                    section.header.is_linked_to_previous = False
                    section.footer.is_linked_to_previous = False
                    section.header.paragraphs[0].clear()
                    section.footer.paragraphs[0].clear()
                    section.header_distance = Cm(0.1)
                    section.footer_distance = Cm(0.1)
                    section.top_margin = Cm(0.5)
                    section.bottom_margin = Cm(0.5)
                    section.left_margin = Cm(0.5)
                    section.right_margin = Cm(0.5)

                for i in range(0, len(sorted_stts), max_per_page):
                    page_stts = sorted_stts[i:i+max_per_page]
                    table = doc.add_table(rows=rows_per_page, cols=cols_per_page)
                    # Set padding trong bảng về 0 để ảnh sát nhau
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    tblCellMar = OxmlElement('w:tblCellMar')
                    for margin in ['top', 'left', 'bottom', 'right']:
                        node = OxmlElement(f'w:{margin}')
                        node.set(qn('w:w'), '0')
                        node.set(qn('w:type'), 'dxa')
                        tblCellMar.append(node)
                    tblPr.append(tblCellMar)
                    table.autofit = False

                    for idx, stt in enumerate(page_stts):
                        if not image_map[stt]["valid"]:
                            continue  # Không chèn ảnh bị lỗi vào Word
                        row = idx // cols_per_page
                        col = idx % cols_per_page
                        cell = table.cell(row, col)

                        # Xóa nội dung cũ
                        for para in cell.paragraphs:
                            para.clear()

                        # Canh trái ảnh bên trái, phải ảnh bên phải
                        alignment = 0 if col == 0 else 2

                        # Tạo bảng con: 1 hàng, 2 cột (STT | Ảnh)
                        inner_table = cell.add_table(rows=1, cols=2)
                        inner_table.autofit = False
                        inner_table.allow_autofit = False

                        # Cấu hình width chính xác
                        inner_table.autofit = False
                        inner_table.allow_autofit = False
                        inner_table.columns[0].width = Cm(0.5)
                        inner_table.columns[1].width = Cm(17.3)  # 7.83 + 0.3 dư

                        # Tổng width của bảng con
                        total_inner_table_width = Cm(9.33)
                        tbl_pr = inner_table._tbl.tblPr

                        tbl_w = OxmlElement('w:tblW')
                        tbl_w.set(qn('w:w'), str(int(total_inner_table_width.cm * 567)))  # 1cm = 567 twips
                        tbl_w.set(qn('w:type'), 'dxa')
                        tbl_pr.append(tbl_w)

                        # Dịch bảng con sang phải chút (nếu cần)
                        tbl_indent = OxmlElement('w:tblInd')
                        tbl_indent.set(qn('w:w'), '150')  # 150 twips = 0.3cm
                        tbl_indent.set(qn('w:type'), 'dxa')
                        tbl_pr.append(tbl_indent)

                        # Cột STT (trái)
                        stt_cell = inner_table.cell(0, 0)
                        stt_cell.width = Cm(1.2)
                        stt_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        stt_para = stt_cell.paragraphs[0]
                        stt_para.paragraph_format.space_after = Pt(0)
                        stt_para.paragraph_format.space_before = Pt(0)
                        stt_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        stt_run = stt_para.add_run(str(stt))
                        stt_run.bold = True
                        stt_run.font.size = Pt(15)


                        # Cột ảnh (phải)
                        img_cell = inner_table.cell(0, 1)
                        img_cell.width = image_width
                        img_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        img_para = img_cell.paragraphs[0]
                        img_para.paragraph_format.space_after = Pt(0)
                        img_para.paragraph_format.space_before = Pt(0)
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        img_run = img_para.add_run()
                        if image_map[stt]["valid"]:
                            img_run.add_picture(image_map[stt]["path"], width=image_width)
                        else:
                            img_run = img_para.add_run("Link không chứa ảnh")
                            img_run.font.size = Pt(11)
                            img_run.bold = True
                            img_run.italic = True

                    # Ngắt trang thủ công sau mỗi bảng (trừ trang cuối)
                    if i + max_per_page < len(sorted_stts):
                        doc.add_page_break()


                # === Xuất file
                base_name = os.path.splitext(sheet_file.name)[0]
                docx_path = f"{base_name}_output.docx"
                doc.save(docx_path)
                st.success("✅ Đã tạo file Word!")

                with open(docx_path, "rb") as f:
                    st.download_button("📥 Tải file Word", f, file_name=docx_path)

                # if export_pdf:
                #     from docx2pdf import convert
                #     pdf_path = f"{base_name}_output.pdf"
                #     convert(docx_path, pdf_path)
                #     with open(pdf_path, "rb") as f:
                #         st.download_button("📥 Tải file PDF", f, file_name=pdf_path)

                # Cleanup
                for item in image_map.values():
                    if item["valid"] and os.path.exists(item["path"]):
                        os.remove(item["path"])


        except Exception as e:
            st.error(f"❌ Lỗi xử lý: {e}")